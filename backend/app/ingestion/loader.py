"""Load supported documents from the local sample document folder.

Supported formats are `.txt`, `.pdf`, and `.docx`. This module extracts text
only; chunking and embeddings are handled in later stages.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader


PROJECT_ROOT = Path(__file__).resolve().parents[3]
DEFAULT_SAMPLE_DOCS_DIR = PROJECT_ROOT / "data" / "sample_docs"
SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


@dataclass(frozen=True)
class LoadedPage:
    """Text extracted from one source page when page numbers are available."""

    page_number: int | None
    text: str


@dataclass(frozen=True)
class LoadedDocument:
    """A supported document loaded from disk."""

    filename: str
    source_path: Path
    source_type: str
    text: str
    pages: list[LoadedPage]


def get_source_type(path: Path) -> str:
    """Return the source type stored in the database for a file."""
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return "sample_txt"
    if suffix == ".pdf":
        return "sample_pdf"
    if suffix == ".docx":
        return "sample_docx"
    raise ValueError(f"Unsupported file type: {path.suffix}")


def load_txt_file(path: Path) -> LoadedDocument:
    """Load a UTF-8 plain text file."""
    text = path.read_text(encoding="utf-8").strip()
    return LoadedDocument(
        filename=path.name,
        source_path=path,
        source_type=get_source_type(path),
        text=text,
        pages=[LoadedPage(page_number=None, text=text)] if text else [],
    )


def load_pdf_file(path: Path) -> LoadedDocument:
    """Extract text from a PDF file using pypdf."""
    reader = PdfReader(str(path))
    pages: list[LoadedPage] = []

    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(LoadedPage(page_number=index, text=text))

    combined_text = "\n\n".join(page.text for page in pages)
    return LoadedDocument(
        filename=path.name,
        source_path=path,
        source_type=get_source_type(path),
        text=combined_text,
        pages=pages,
    )


def load_docx_file(path: Path) -> LoadedDocument:
    """Extract paragraph text from a Word `.docx` file."""
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraph for paragraph in paragraphs if paragraph).strip()

    return LoadedDocument(
        filename=path.name,
        source_path=path,
        source_type=get_source_type(path),
        text=text,
        pages=[LoadedPage(page_number=None, text=text)] if text else [],
    )


def load_document(path: Path) -> LoadedDocument:
    """Load one supported document file."""
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return load_txt_file(path)
    if suffix == ".pdf":
        return load_pdf_file(path)
    if suffix == ".docx":
        return load_docx_file(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def load_documents(directory: Path = DEFAULT_SAMPLE_DOCS_DIR) -> list[LoadedDocument]:
    """Load all supported document files from a directory.

    Args:
        directory: Folder containing sample `.txt`, `.pdf`, or `.docx` files.

    Returns:
        A list of loaded documents sorted by filename.
    """
    if not directory.exists():
        raise FileNotFoundError(f"Sample document folder does not exist: {directory}")

    documents: list[LoadedDocument] = []
    for path in sorted(directory.iterdir()):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        document = load_document(path)
        if not document.text:
            print(f"Skipping empty document: {path.name}")
            continue

        documents.append(document)

    return documents


def load_txt_files(directory: Path = DEFAULT_SAMPLE_DOCS_DIR) -> list[LoadedDocument]:
    """Backward-compatible wrapper for older stages.

    New code should use `load_documents`.
    """
    if not directory.exists():
        raise FileNotFoundError(f"Sample document folder does not exist: {directory}")
    return [load_txt_file(path) for path in sorted(directory.glob("*.txt")) if path.is_file()]
